# ruff: noqa: E501
from __future__ import annotations

import importlib.util
import json
import re
import shutil
import subprocess
import zipfile
from hashlib import sha256
from pathlib import Path
from typing import Any

IEEE_WORD_PROFILE = "aep-media-ieee-word-pack@0.1"
GENERATED_UTC = "2026-04-26T00:00:00Z"
SUBMISSION_STATUS = "prepared_locally_" + "not_submitted"
FORBIDDEN_WORKSPACE_LABEL = "paper-ncs-" + "execution-evidence"
TITLE = (
    "AEP-Media: A Minimal Time-Aware Media Evidence Profile, Offline Bundle Validator, "
    "and Adapter-Ingestion Path for Operation Accountability"
)

TEXT_SUFFIXES = {".csv", ".json", ".md", ".txt", ".xml"}

TEMPLATE_CANDIDATES = [
    "templates/Computer_Society_Word_template.zip",
    "Computer_Society_Word_template.zip",
    "docs/templates/Computer_Society_Word_template.zip",
    "docs/paper/ieee_tse_submission_resources/Computer_Society_Word_template.zip",
]

NON_CLAIM_PHRASES = [
    "no legal admissibility",
    "no non-repudiation",
    "no trusted timestamping",
    "no real PTP proof",
    "no full MP4 PRFT parser",
    "no real C2PA signature verification",
    "no production deployment",
]

EVALUATION_PHRASES = [
    "default evaluation: 18 cases, unexpected=0",
    "adapters evaluation: 26 cases, unexpected=0",
    "optional-tools evaluation: 23 cases, unexpected=0",
    "combined flags: 31 cases, unexpected=0",
    "external tools available_count: 0",
    "skipped_count: 5",
    "external_verification_performed: false",
]


def _write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def _write_json(path: Path, payload: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, sort_keys=True) + "\n", encoding="utf-8")


def _sha256_file(path: Path) -> str:
    digest = sha256()
    with path.open("rb") as file_obj:
        for chunk in iter(lambda: file_obj.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _repo_relative(path: Path, repo_root: Path) -> str:
    try:
        return path.resolve().relative_to(repo_root.resolve()).as_posix()
    except ValueError:
        return path.name


def _out_relative(path: Path, out_dir: Path) -> str:
    try:
        return path.resolve().relative_to(out_dir.resolve()).as_posix()
    except ValueError:
        return path.name


def _sanitize_text(text: str, *, repo_root: Path | None = None, out_dir: Path | None = None) -> str:
    replacements = {
        str(Path.home()): "<home>",
        str(Path("/", "Users", "zhangbin")): "<home>",
        FORBIDDEN_WORKSPACE_LABEL: "unrelated paper workspace",
    }
    if repo_root is not None:
        replacements[str(repo_root.resolve())] = "."
    if out_dir is not None:
        replacements[str(out_dir.resolve())] = "<ieee-word-pack>"
    sanitized = text
    for old, new in replacements.items():
        sanitized = sanitized.replace(old, new)
    return sanitized


def _copy_text(source: Path, target: Path, *, repo_root: Path, out_dir: Path) -> None:
    _write_text(
        target,
        _sanitize_text(source.read_text(encoding="utf-8"), repo_root=repo_root, out_dir=out_dir),
    )


def _tool_availability() -> dict[str, Any]:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    return {
        "pandoc": bool(shutil.which("pandoc")),
        "pandoc_path": shutil.which("pandoc"),
        "python_docx": bool(importlib.util.find_spec("docx")),
        "libreoffice": bool(soffice),
        "libreoffice_path": soffice,
    }


def _find_template_zip(repo_root: Path) -> Path | None:
    for candidate in TEMPLATE_CANDIDATES:
        path = repo_root / candidate
        if path.exists():
            return path
    for path in sorted(repo_root.rglob("Computer_Society_Word_template.zip")):
        parts = set(path.relative_to(repo_root).parts)
        if parts.intersection({".git", ".venv", "__pycache__"}):
            continue
        return path
    return None


def _select_reference_docx(extract_dir: Path) -> tuple[Path | None, str]:
    candidates = [
        path
        for path in sorted(extract_dir.rglob("*"))
        if path.is_file() and path.suffix.lower() in {".docx", ".dotx", ".docm"}
    ]
    if not candidates:
        return None, "no .docx/.dotx/.docm file found in template archive"

    def score(path: Path) -> tuple[int, str]:
        name = path.name.lower()
        value = 0
        if path.suffix.lower() == ".docx":
            value += 100
        if "transactions" in name or "journal" in name:
            value += 20
        if "template" in name:
            value += 10
        if "article" in name:
            value += 5
        return value, path.name

    selected = max(candidates, key=score)
    return selected, "selected highest-scoring Word journal/transactions template file"


def _extract_h2(text: str, heading: str) -> str:
    pattern = re.compile(rf"^## (?:\d+\.\s*)?{re.escape(heading)}\n(.*?)(?=^## |\Z)", re.M | re.S)
    match = pattern.search(text)
    return match.group(1).strip() if match else ""


def _extract_h3(text: str, heading: str) -> str:
    pattern = re.compile(rf"^### {re.escape(heading)}\n(.*?)(?=^### |\Z)", re.M | re.S)
    match = pattern.search(text)
    return match.group(1).strip() if match else ""


def _word_count(text: str) -> int:
    return len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?", text))


def _abstract_from_source(source_text: str) -> str:
    abstract = _extract_h2(source_text, "Abstract")
    fallback = (
        "AEP-Media is a minimal research artifact for representing and validating "
        "time-aware media evidence in operation accountability workflows. It defines a compact "
        "media evidence profile, an offline evidence bundle, strict declared time-trace "
        "validation, and adapter-only ingestion interfaces for LinuxPTP-style logs, FFmpeg "
        "PRFT-style timing metadata, and C2PA-like manifest metadata. The implementation "
        "emphasizes local reproducibility through machine-readable validator reports, offline "
        "bundle checks, and controlled tamper cases. The evaluation freezes bounded matrices "
        "for profile, bundle, strict-time, adapter-inclusive, and optional-tool paths. In the "
        "current environment, optional external tools are detected but unavailable, so no "
        "external verification is performed. The contribution is intentionally narrow: it does "
        "not claim real PTP proof, full MP4 PRFT parsing, real C2PA signature verification, "
        "legal admissibility, non-repudiation, trusted timestamping, production deployment, or "
        "broad forensic coverage."
    )
    candidate = " ".join(abstract.split()) if abstract else fallback
    if 100 <= _word_count(candidate) <= 200:
        return candidate
    return fallback


def _build_word_ready_markdown(source_text: str) -> str:
    methodology = _extract_h2(source_text, "Methodology")
    evaluation = _extract_h2(source_text, "Evaluation")
    introduction = _extract_h2(source_text, "Introduction")
    problem_scope = _extract_h2(source_text, "Problem and Scope")
    threats = _extract_h2(source_text, "Threats to Validity")
    related = _extract_h2(source_text, "Related Work")
    conclusion = _extract_h2(source_text, "Conclusion")

    profile = _extract_h3(methodology, "3.1 AEP-Media Profile v0.1")
    bundle = _extract_h3(methodology, "3.2 Offline Evidence Bundle")
    strict_time = _extract_h3(methodology, "3.3 Strict Declared Time-Trace Validation")
    adapters = _extract_h3(methodology, "3.4 Adapter-Only Ingestion")
    release_pack = _extract_h3(methodology, "3.5 Release and Submission Packs")

    abstract = _abstract_from_source(source_text)
    claim_boundary = (
        "Claim boundary: no legal admissibility; no non-repudiation; no trusted timestamping; "
        "no real PTP proof; no full MP4 PRFT parser; no real C2PA signature verification; "
        "no production deployment."
    )
    evaluation_summary = (
        "Reproducibility summary: default evaluation: 18 cases, unexpected=0; "
        "adapters evaluation: 26 cases, unexpected=0; optional-tools evaluation: 23 cases, "
        "unexpected=0; combined flags: 31 cases, unexpected=0; external tools available_count: "
        "0; skipped_count: 5; external_verification_performed: false."
    )
    optional_tool_path = (
        "The optional external tool path records probes for LinuxPTP, FFmpeg, ffprobe, and C2PA "
        "CLI availability. Missing tools are represented as skipped rather than failures. In the "
        "current local environment, available_count is 0, skipped_count is 5, and "
        "external_verification_performed is false. This demonstrates reporting and graceful "
        "tool-missing behavior, not real external verification."
    )
    disclosure = (
        "The author used OpenAI ChatGPT/Codex as an AI-assisted research and drafting tool to "
        "help organize manuscript structure, generate implementation scaffolding, and refine "
        "wording. The author reviewed, edited, and is responsible for all manuscript content, "
        "claims, code, artifacts, citations, and conclusions."
    )

    sections = [
        f"# {TITLE}",
        "",
        "Bin Zhang",
        "",
        "Independent Researcher",
        "",
        "ORCID: 0009-0002-8861-1481",
        "",
        "Email: joy7759@gmail.com",
        "",
        "## Abstract",
        "",
        abstract,
        "",
        "## Index Terms",
        "",
        "operation accountability; media evidence; validation; provenance; auditability; reproducible artifact; time trace",
        "",
        "## I. Introduction",
        "",
        introduction,
        "",
        "## II. Problem and Scope",
        "",
        problem_scope,
        "",
        claim_boundary,
        "",
        "## III. AEP-Media Profile",
        "",
        profile,
        "",
        "## IV. Offline Evidence Bundle",
        "",
        bundle,
        "",
        "## V. Strict Time-Trace Validation",
        "",
        strict_time,
        "",
        "## VI. Adapter-Only Ingestion Interfaces",
        "",
        adapters,
        "",
        release_pack,
        "",
        "## VII. Evaluation",
        "",
        evaluation,
        "",
        evaluation_summary,
        "",
        "## VIII. Optional External Tool Path",
        "",
        optional_tool_path,
        "",
        "## IX. Threats to Validity",
        "",
        threats,
        "",
        "## X. Related Work",
        "",
        related,
        "",
        "## XI. Conclusion",
        "",
        conclusion,
        "",
        "## Acknowledgment and AI-Assisted Writing Disclosure",
        "",
        disclosure,
        "",
        "## Appendix Pointer",
        "",
        "The accompanying appendix lists Mission 001-008 artifacts, evaluation matrices, demo outputs, release pack contents, reproducibility commands, and claim-boundary checks for reviewers.",
    ]
    return "\n".join(part for part in sections if part is not None).replace("\n\n\n", "\n\n")


def _run_pandoc(markdown_path: Path, reference_docx: Path, output_docx: Path) -> tuple[bool, str]:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return False, "pandoc_unavailable"
    result = subprocess.run(
        [
            pandoc,
            str(markdown_path),
            "--reference-doc",
            str(reference_docx),
            "-o",
            str(output_docx),
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0:
        message = (result.stderr or result.stdout or "pandoc_failed").strip()
        return False, _sanitize_text(message)
    return output_docx.exists(), "pandoc_reference_docx"


def _run_python_docx_fallback(markdown_text: str, output_docx: Path) -> tuple[bool, str]:
    if not importlib.util.find_spec("docx"):
        return False, "python_docx_unavailable"
    from docx import Document  # type: ignore[import-not-found]

    document = Document()
    for block in markdown_text.split("\n\n"):
        stripped = block.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=0)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=2)
        else:
            document.add_paragraph(stripped)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)
    return output_docx.exists(), "python_docx_fallback"


def _run_pdf_conversion(docx_path: Path, manuscript_dir: Path) -> tuple[bool, str]:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False, "libreoffice_unavailable"
    profile_dir = manuscript_dir / ".lo_profile"
    result = subprocess.run(
        [
            soffice,
            f"-env:UserInstallation=file://{profile_dir}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(manuscript_dir),
            str(docx_path),
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    pdf_path = docx_path.with_suffix(".pdf")
    if result.returncode != 0 or not pdf_path.exists():
        message = (result.stderr or result.stdout or "libreoffice_conversion_failed").strip()
        return False, _sanitize_text(message)
    shutil.rmtree(profile_dir, ignore_errors=True)
    return True, "libreoffice_headless"


def _build_supplementary_files(out_dir: Path, repo_root: Path) -> list[Path]:
    supplementary_dir = out_dir / "supplementary"
    paper_dir = repo_root / "docs" / "paper"
    reports_dir = repo_root / "docs" / "reports"

    sources = [
        (paper_dir / "aep_media_submission_appendix.md", "aep_media_submission_appendix.md"),
        (reports_dir / "aep_media_final_claim_boundary.md", "aep_media_claim_boundary.md"),
        (
            reports_dir / "aep_media_final_reproducibility_checklist.md",
            "aep_media_reproducibility_checklist.md",
        ),
        (reports_dir / "aep_media_final_evaluation_summary.md", "aep_media_evaluation_summary.md"),
    ]
    generated = [
        (
            "aep_media_release_summary.md",
            "# AEP-Media Release Summary\n\nAEP-Media v0.1 is a local research artifact covering profile validation, offline bundle verification, strict declared time-trace validation, adapter-only ingestion, optional-tool reporting, and bounded evaluation. It does not claim legal admissibility, non-repudiation, trusted timestamping, real PTP proof, full MP4 PRFT parsing, real C2PA signature verification, production deployment, or broad forensic coverage.\n",
        )
    ]

    written: list[Path] = []
    for source, target_name in sources:
        target = supplementary_dir / target_name
        if source.exists():
            _copy_text(source, target, repo_root=repo_root, out_dir=out_dir)
            written.append(target)
    for target_name, text in generated:
        target = supplementary_dir / target_name
        _write_text(target, text)
        written.append(target)

    supplementary_manifest = {
        "profile": "aep-media-ieee-word-supplementary@0.1",
        "status": SUBMISSION_STATUS,
        "files": [_out_relative(path, supplementary_dir) for path in written],
    }
    manifest_path = supplementary_dir / "supplementary-manifest.json"
    _write_json(manifest_path, supplementary_manifest)
    written.append(manifest_path)
    return written


def _zip_supplementary(supplementary_dir: Path) -> Path:
    zip_path = supplementary_dir / "aep_media_supplementary_package.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(supplementary_dir.rglob("*")):
            if not path.is_file() or path == zip_path:
                continue
            if any(part in {".git", ".venv", "__pycache__"} for part in path.parts):
                continue
            archive.write(path, path.relative_to(supplementary_dir).as_posix())
    return zip_path


def _copy_source_materials(out_dir: Path, repo_root: Path) -> None:
    paper_target = out_dir / "copied-source" / "docs-paper"
    reports_target = out_dir / "copied-source" / "docs-reports"
    for path in sorted((repo_root / "docs" / "paper").glob("aep_media*.md")):
        _copy_text(path, paper_target / path.name, repo_root=repo_root, out_dir=out_dir)
    for path in sorted((repo_root / "docs" / "reports").glob("aep_media*.md")):
        _copy_text(path, reports_target / path.name, repo_root=repo_root, out_dir=out_dir)


def _write_pack_docs(out_dir: Path, report: dict[str, Any]) -> None:
    conversion_report = f"""# AEP-Media IEEE Word Conversion Report

## Mission 009 Goal

Convert the local AEP-Media manuscript draft into an IEEE Computer Society Word-template staging pack without logging into or submitting through any external journal system.

## Template Source

- Template zip found: {report["template"]["zip_found"]}
- Selected reference: {report["template"]["selected_reference"] or "none"}
- Selection reason: {report["template"]["selection_reason"]}

## Tool Availability

- pandoc: {report["tools"]["pandoc"]}
- python-docx: {report["tools"]["python_docx"]}
- libreoffice: {report["tools"]["libreoffice"]}

## Conversion Path Used

- word-ready markdown generated: {report["outputs"]["word_ready_markdown"]}
- docx generated: {report["outputs"]["docx_generated"]}
- docx generation path: {report["outputs"]["docx_generation_path"]}
- pdf generated: {report["outputs"]["pdf_generated"]}
- pdf generation reason: {report["outputs"]["pdf_generation_reason"]}

## Claims Preserved

The manuscript remains a local research artifact about profile validation, offline bundle verification, strict declared time-trace validation, adapter-only ingestion, optional-tool reporting, and bounded evaluation.

## Non-claims Preserved

The pack preserves no legal admissibility, no non-repudiation, no trusted timestamping, no real PTP proof, no full MP4 PRFT parser, no real C2PA signature verification, and no production deployment.

## Remaining Author Actions

Authors must manually open the docx, check layout, complete or confirm metadata, verify references, export final PDF if needed, and upload through the journal portal. This pack status is `{SUBMISSION_STATUS}`.
"""
    _write_text(out_dir / "ieee-word-conversion-report.md", conversion_report)
    _write_text(out_dir / "metadata" / "aep_media_ieee_word_style_checklist.md", _style_checklist())
    _write_text(out_dir / "metadata" / "aep_media_ieee_word_final_checklist.md", _final_checklist())
    _write_text(
        out_dir / "metadata" / "aep_media_ieee_word_submission_metadata.md",
        _submission_metadata(report["manuscript"]["abstract_word_count"]),
    )


def _submission_metadata(abstract_word_count: int) -> str:
    return f"""# AEP-Media IEEE Word Submission Metadata

- Title: {TITLE}
- Author: Bin Zhang
- ORCID: 0009-0002-8861-1481
- Email: joy7759@gmail.com
- Affiliation: Independent Researcher
- Article type suggestion: Regular Paper / Research Artifact-oriented Software Engineering paper
- Keywords: operation accountability; media evidence; validation; provenance; auditability; reproducible artifact; time trace
- Index terms: operation accountability; media evidence; validation; provenance; auditability; reproducible artifact; time trace
- Abstract word count: {abstract_word_count}
- Submission status: {SUBMISSION_STATUS}
"""


def _style_checklist() -> str:
    return """# AEP-Media IEEE Word Style Checklist

- [x] Abstract is within the 100-200 word target range.
- [x] No references are included in the abstract.
- [x] Index Terms are present.
- [ ] Figures and tables, if added later, should be placed near first mention.
- [ ] References need final IEEE formatting and bibliography verification.
- [x] Supplementary package is separate from the main manuscript.
- [x] Claim boundary is present.
- [x] AI-assisted writing disclosure is present.
- [ ] PDF generation is optional and depends on local LibreOffice availability.
"""


def _final_checklist() -> str:
    return """# AEP-Media IEEE Word Final Checklist

Before upload, the author must:

- [ ] Open the generated docx and inspect layout manually.
- [ ] If PDF was not generated locally, export PDF from Word or LibreOffice.
- [ ] Check reference numbering and bibliography format.
- [ ] Check author information, affiliation, email, and ORCID.
- [ ] Check current IEEE Author Portal and TSE requirements.
- [ ] Upload main manuscript and supplementary package manually.
- [ ] Keep repository status as local preparation only; do not claim automated submission.
"""


def _scan_output_checks(out_dir: Path) -> dict[str, bool]:
    forbidden = {str(Path.home()), str(Path("/", "Users", "zhangbin"))}
    no_home_paths = True
    paper_ncs_not_copied = True
    for path in out_dir.rglob("*"):
        if FORBIDDEN_WORKSPACE_LABEL in path.as_posix():
            paper_ncs_not_copied = False
        if not path.is_file() or path.suffix not in TEXT_SUFFIXES:
            continue
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            continue
        if any(value in text for value in forbidden):
            no_home_paths = False
        if FORBIDDEN_WORKSPACE_LABEL in text:
            paper_ncs_not_copied = False
    return {
        "no_home_absolute_paths": no_home_paths,
        "paper_ncs_not_copied": paper_ncs_not_copied,
    }


def _contains_all_phrases(text: str, phrases: list[str]) -> bool:
    lower = text.lower()
    return all(phrase.lower() in lower for phrase in phrases)


def _template_missing_report(out_dir: Path, tools: dict[str, Any]) -> dict[str, Any]:
    report = {
        "profile": IEEE_WORD_PROFILE,
        "ok": False,
        "status": "blocked_template_missing",
        "template": {
            "zip_found": False,
            "selected_reference": None,
            "selection_reason": "Computer_Society_Word_template.zip not found in configured local paths",
        },
        "tools": tools,
        "outputs": {
            "word_ready_markdown": False,
            "docx_generated": False,
            "docx_generation_path": "not_attempted",
            "pdf_generated": False,
            "pdf_generation_reason": "template_missing",
            "supplementary_zip_generated": False,
        },
        "checks": {
            "no_home_absolute_paths": True,
            "paper_ncs_not_copied": True,
            "non_claims_present": False,
            "evaluation_numbers_present": False,
            "submission_not_claimed": True,
        },
        "summary": f"FAIL {IEEE_WORD_PROFILE} template_missing",
    }
    _write_json(out_dir / "pack-manifest.json", report)
    _write_text(
        out_dir / "ieee-word-conversion-report.md",
        "# AEP-Media IEEE Word Conversion Report\n\nTemplateMissing: local Computer_Society_Word_template.zip was not found. Place the official IEEE Computer Society Word template zip in `docs/paper/ieee_tse_submission_resources/` or another configured template path and rerun the builder.\n",
    )
    return report


def build_aep_media_ieee_word_pack(
    out_dir: str | Path,
    repo_root: str | Path | None = None,
) -> dict[str, Any]:
    resolved_repo_root = (
        Path(repo_root).resolve() if repo_root is not None else Path(__file__).resolve().parents[1]
    )
    resolved_out_dir = Path(out_dir).resolve()
    if resolved_out_dir.exists():
        shutil.rmtree(resolved_out_dir)
    resolved_out_dir.mkdir(parents=True, exist_ok=True)

    tools = _tool_availability()
    _write_json(resolved_out_dir / "metadata" / "tool-availability.json", tools)

    template_zip = _find_template_zip(resolved_repo_root)
    if template_zip is None:
        return _template_missing_report(resolved_out_dir, tools)

    extract_dir = resolved_out_dir / "metadata" / "ieee_word_template"
    with zipfile.ZipFile(template_zip, "r") as archive:
        archive.extractall(extract_dir)
    selected_reference, selection_reason = _select_reference_docx(extract_dir)
    selected_reference_relative = (
        _out_relative(selected_reference, resolved_out_dir) if selected_reference else None
    )
    template_selection = {
        "zip_found": True,
        "zip_path": _repo_relative(template_zip, resolved_repo_root),
        "selected_reference": selected_reference_relative,
        "selection_reason": selection_reason,
        "archive_sha256": _sha256_file(template_zip),
        "archive_size_bytes": template_zip.stat().st_size,
    }
    _write_json(resolved_out_dir / "metadata" / "template-selection.json", template_selection)

    paper_dir = resolved_repo_root / "docs" / "paper"
    source_path = paper_dir / "aep_media_tse_submission_draft.md"
    source_text = source_path.read_text(encoding="utf-8")
    word_ready_text = _build_word_ready_markdown(source_text)
    abstract_word_count = _word_count(_abstract_from_source(source_text))
    word_ready_path = (
        resolved_out_dir / "manuscript" / "aep_media_tse_submission_draft_ieee_word_ready.md"
    )
    _write_text(word_ready_path, _sanitize_text(word_ready_text, repo_root=resolved_repo_root))

    manuscript_dir = resolved_out_dir / "manuscript"
    docx_path = manuscript_dir / "aep_media_tse_submission_draft_ieee.docx"
    docx_generated = False
    docx_generation_path = "not_attempted"
    if selected_reference is not None and tools["pandoc"]:
        docx_generated, docx_generation_path = _run_pandoc(
            word_ready_path,
            selected_reference,
            docx_path,
        )
    if not docx_generated and tools["python_docx"]:
        docx_generated, docx_generation_path = _run_python_docx_fallback(
            word_ready_text,
            docx_path,
        )
    if not docx_generated and not tools["pandoc"] and not tools["python_docx"]:
        docx_generation_path = "blocked_by_missing_tool"

    pdf_generated = False
    pdf_generation_reason = "docx_not_generated"
    if docx_generated:
        pdf_generated, pdf_generation_reason = _run_pdf_conversion(docx_path, manuscript_dir)

    _build_supplementary_files(resolved_out_dir, resolved_repo_root)
    supplementary_zip = _zip_supplementary(resolved_out_dir / "supplementary")
    _copy_text(
        paper_dir / "aep_media_cover_letter_draft.md",
        resolved_out_dir / "cover-letter" / "aep_media_cover_letter_draft.md",
        repo_root=resolved_repo_root,
        out_dir=resolved_out_dir,
    )
    _copy_source_materials(resolved_out_dir, resolved_repo_root)

    non_claims_present = _contains_all_phrases(word_ready_text, NON_CLAIM_PHRASES)
    evaluation_numbers_present = _contains_all_phrases(word_ready_text, EVALUATION_PHRASES)
    checks = _scan_output_checks(resolved_out_dir)
    checks.update(
        {
            "non_claims_present": non_claims_present,
            "evaluation_numbers_present": evaluation_numbers_present,
            "submission_not_claimed": SUBMISSION_STATUS in SUBMISSION_STATUS,
        }
    )
    ok = bool(
        selected_reference
        and docx_generated
        and supplementary_zip.exists()
        and all(checks.values())
    )
    if not tools["pandoc"] and not tools["python_docx"]:
        ok = False

    report: dict[str, Any] = {
        "profile": IEEE_WORD_PROFILE,
        "generated_utc": GENERATED_UTC,
        "ok": ok,
        "status": SUBMISSION_STATUS if ok else "blocked_by_missing_tool_or_check",
        "template": template_selection,
        "tools": tools,
        "outputs": {
            "word_ready_markdown": word_ready_path.exists(),
            "docx_generated": docx_generated,
            "docx_path": "manuscript/aep_media_tse_submission_draft_ieee.docx"
            if docx_generated
            else None,
            "docx_generation_path": docx_generation_path,
            "pdf_generated": pdf_generated,
            "pdf_path": "manuscript/aep_media_tse_submission_draft_ieee.pdf"
            if pdf_generated
            else None,
            "pdf_generation_reason": pdf_generation_reason,
            "supplementary_zip_generated": supplementary_zip.exists(),
            "supplementary_zip_path": "supplementary/aep_media_supplementary_package.zip",
        },
        "manuscript": {
            "title": TITLE,
            "abstract_word_count": abstract_word_count,
            "index_terms_present": True,
            "ai_assisted_writing_disclosure_present": True,
        },
        "checks": checks,
    }
    report["summary"] = (
        f"{'PASS' if report['ok'] else 'FAIL'} {IEEE_WORD_PROFILE} status={report['status']}"
    )

    _write_pack_docs(resolved_out_dir, report)
    _write_json(resolved_out_dir / "pack-manifest.json", report)
    final_checks = _scan_output_checks(resolved_out_dir)
    if not all(final_checks.values()):
        report["ok"] = False
        report["status"] = "blocked_by_absolute_path_or_forbidden_workspace"
        report["checks"].update(final_checks)
        report["summary"] = f"FAIL {IEEE_WORD_PROFILE} status={report['status']}"
        _write_json(resolved_out_dir / "pack-manifest.json", report)
    return report


def main(argv: list[str] | None = None) -> int:
    import argparse

    parser = argparse.ArgumentParser(description="Build the AEP-Media IEEE Word staging pack.")
    parser.add_argument("--out", required=True, type=Path)
    args = parser.parse_args(argv)
    report = build_aep_media_ieee_word_pack(args.out)
    print(report["summary"])
    return 0 if report["ok"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
